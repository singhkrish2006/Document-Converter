from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pdfplumber

def word_to_pdf(word_file, pdf_file):
    doc = Document(word_file)
    pdf = canvas.Canvas(pdf_file, pagesize=letter)
    width, height = letter
    y_position = height - 50  

    for para in doc.paragraphs:
        pdf.drawString(50, y_position, para.text)
        y_position -= 20  

        if y_position < 50:
            pdf.showPage()
            y_position = height - 50
    
    pdf.save()
    print(f"Successfully converted {word_file} to {pdf_file}")

def pdf_to_word(pdf_file, word_file):
    doc = Document()
    
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
    
    doc.save(word_file)
    print(f"Successfully converted {pdf_file} to {word_file}")

